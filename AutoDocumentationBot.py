import telebot
import os
import tempfile
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell

token = '8582480612:AAGL3Dg6hYJLKO09jmxROrQ-E4l7yl3HKP4'
bot = telebot.TeleBot(token)

user_states = {}

#приветствие и навигация
@bot.message_handler(commands=['start'])
def start_mes(message):
    bot.send_message(message.chat.id, "Здравствуйте!\nЯ - бот, который автоматизирует и упрощает работу с документами!\nЯ был разработан студенткой группы ИС-22-1 "
    "специально с целью упрощения монотонной работы, которая забирает огромную кучу времени у преподавателей учреждения, поэтому и смог увидеть этот удивительный цифровой мир, чтобы облегчить работу людям!\nПожалуйста, "
    "выберите одну из следующих команд, чтобы начать работу со мной!\n/help - Помощь, если у вас "
    "возникли трудности\n/tags - Заметки, где и какие тэги нужно использовать для заполнения\n/fill_docx - Заполнить документ!")
    bot.send_message(message.chat.id, "Для начала работы, ознакомьтесь с тэгами для заполнения через команду /tags и начинайте работу с вашими документами через команду /fill_docx!")

@bot.message_handler(commands=['help'])
def help_mes(message):
    bot.send_message(message.chat.id, "<b>Как меня использовать?</b>\n<b>Шаг 1.</b> Нажмите на команду /tags - она даст вам полное представление, какие тэги и в каких документах они используются.\n<b>Шаг 2.</b> "
    "После ознакомления с тэгами, воспользуйтесь командой /fill_docx, чтобы загрузить документ, который хотите заполнить с моей помощью!\n<b>Шаг 3.</b> "
    "Загрузите шаблон документа. Если у вас нет нужного  документа, можете скачать его здесь: https://drive.google.com/drive/folders/1SNSlid6au3tRX-YGSklDyoUYeqWxB1qh?usp=sharing\n<b>Шаг 4.</b> "
    "Заполните нужные для документа тэги и отправьте их мне!\n<b>Шаг 5.</b> "
    "Теперь стоит немного подождать! Я обработаю ваш документ и отправлю вам заполненный!", parse_mode='HTML')

@bot.message_handler(commands=['tags'])
def tags_mes(message):
    bot.send_message(message.chat.id, "Вот какие тэги используются в документах:\n" \
    "<b>ИС УП.01.02</b>\n" \
    "Аттестационный лист – <code>code_group: Код группы, student_name: ФИО студента, num_course: Номер курса, " \
    "num_sem: Номер семестра, start_pract: Дата начала учебной практики, end_pract: Дата конца учебной практики, grade: Оценка</code>\n" \
    "Зачетная ведомость – <code>num: Номер зачетной ведомости, code_group: Код группы, num_course: Номер курса, num_sem: Номер семестра, " \
    "start_pract: Дата начала учебной практики, end_pract: Дата конца учебной практики, sn1-sn26: ФИО студента, date: Дата проведения зачета</code>\n" \
    #"**Экзаменационная ведомость** - code_group: Код группы, name_group: Наименование группы, student_name: ФИО студента\n"
    "<b>Аттестационный лист</b> - code_group: Код группы, student_name: ФИО студента, start_pract: Дата начала практики, end_pract: Дата конца практики, grade: Оценка\n"
    "<b>Характеристика на студента</b> - num_course: Номер курса, code_group: Код группы, name_spec: Название специальности, student_name: ФИО студента, "
    "date_birth: Дата рождения студента, adress_student: Адрес проживания студента, fn_student: Имя студента, date_enroll: Дата зачисления на курс, date: Дата заполнения документа", parse_mode='HTML')

#обработка документа
@bot.message_handler(commands=['fill_docx'])
def filldocx(message):
    user_id = message.from_user.id
    user_states[user_id] = 'waiting_document'
    bot.reply_to(message, "Отлично, жду ваш документ!")

@bot.message_handler(content_types=['document'])
def handle_document(message):
    user_id = message.from_user.id
    if user_states.get(user_id) != 'waiting_document':
        return
    
    file_info = bot.get_file(message.document.file_id)
    downloaded_file = bot.download_file(file_info.file_path)

    file_name = message.document.file_name
    if not file_name.lower().endswith('.docx'):
        bot.reply_to(message, "Вы прислали не тот тип документа! Пожалуйста, отправьте документ с расширением .docx!")
        del user_states[user_id]
        return
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
        temp_file.write(downloaded_file)
        temp_file_path = temp_file.name
    
    user_states[user_id] = { 'state': 'waiting_data', 'file_path': temp_file_path }
    bot.send_message(message.chat.id, "С документом все отлично! Пожалуйста, отправьте мне данные в формате - ключ: значение.\n"
    "<b>ВАЖНО!</b>\nДанные нужно заполнять все  и сразу! Например, num_course: 4, code_group: ИС-22-1 и так далее, в зависимости от документа, который вы хотите заполнить", parse_mode='HTML')

#заполнение документа
@bot.message_handler(func=lambda message: user_states.get(message.from_user.id, {}).get('state') == 'waiting_data')
def handle_data(message):
    user_id = message.from_user.id
    state_info = user_states[user_id]
    file_path = state_info['file_path']

    try:
        data_text = message.text
        data_dict = {}
        for pair in data_text.split(','):
            if ':' in pair:
                key, value = pair.split(':', 1)
                data_dict[key.strip()] = value.strip()
        if not data_dict:
            bot.reply_to(message, "Не могу распознать данные. Пожалуйста, используйте формат <b>ключ: значение</b>", parse_mode='HTML')
            return
        doc = Document(file_path)
        def replace_text(element_text, data): #заполнение всех плейсхолдеров в документе
            for key, value in data.items():
                ph = f'{{{{{key}}}}}'
                element_text = element_text.replace(ph, value)
            return element_text
        for paragraph in doc.paragraphs:
            paragraph.text = replace_text(paragraph.text, data_dict)
        def process_tables(tables, data):
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.text = replace_text(paragraph.text, data_dict)

                        process_tables(cell.tables, data)
        process_tables(doc.tables, data_dict)
        filled_file_path = file_path.replace('.docx', '_filled.docx')
        doc.save(filled_file_path)

        with open(filled_file_path, 'rb') as f:
            bot.send_document(user_id, f, caption="Вот ваш заполненный документ!")
        
        os.remove(file_path)
        os.remove(filled_file_path)
        del user_states[user_id]

    except Exception as e:
        bot.reply_to(message, f'Возникла ошибка при  обработке: {str(e)}. Попробуйте снова с командой /fill_docx')
        if os.path.exists(file_path):
            os.remove(file_path)
        del user_states[user_id]

bot.polling() #инициализация